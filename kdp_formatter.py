import json
import math
from pathlib import Path
from datetime import datetime

try:
    from docx import Document # pyre-ignore[21]
    from docx.shared import Inches, Pt # pyre-ignore[21]
    from docx.enum.text import WD_ALIGN_PARAGRAPH # pyre-ignore[21]
    from docx.oxml import OxmlElement # pyre-ignore[21]
    from docx.oxml.ns import qn # pyre-ignore[21]
except ImportError:
    raise ImportError("python-docx is not installed. Please pip install python-docx")

class KDPFormatter:
    """Formats Markdown/Text books into KDP compliant .docx manuscripts."""

    def __init__(self, book_dir: Path):
        self.book_dir = book_dir
        self.metadata = self._load_metadata()

    def _load_metadata(self) -> dict:
        meta_path = self.book_dir / "metadata.json"
        if meta_path.exists():
            with open(meta_path, "r", encoding="utf-8") as f:
                return json.load(f)
        return {}

    def _add_hyperlink(self, paragraph, text, target):
        """Adds a clickable hyperlink within a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(target, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Style the link
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF') # Blue
        rPr.append(color)

        new_run.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
        return hyperlink

    def _add_bookmark(self, paragraph, name):
        """Adds a bookmark at the start of a paragraph."""
        tag = paragraph._p
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), name)
        tag.insert(0, start)
        
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        tag.append(end)

    def _calculate_gutter(self, page_count: int) -> float:
        """
        Dynamic gutter calculation per Amazon KDP specs:
        - 24-150 pages: 0.375"
        - 151-400 pages: 0.75"
        - 401-600 pages: 0.875"
        - 600+ pages: 1.0"
        """
        if page_count < 151:
            return 0.375
        elif page_count < 401:
            return 0.75
        elif page_count < 601:
            return 0.875
        else:
            return 1.0

    def format_for_kdp(self) -> Path:
        """Creates the formal 6x9 compliant DOCX."""
        txt_path = self.book_dir / "book.txt"
        
        if not txt_path.exists():
            raise FileNotFoundError(f"Source text file not found at {txt_path}")

        # Basic KDP 6x9 Trim Size logic
        estimated_pages = self.metadata.get("estimated_pages", 200)
        gutter_inches = self._calculate_gutter(estimated_pages)

        doc = Document()

        # Define styles
        style = doc.styles['Normal']
        font = style.font
        import config # pyre-ignore[21]
        font.name = getattr(config, "KDP_FONT", "Garamond")
        font.size = Pt(11)
        # Leading (line spacing)
        style.paragraph_format.line_spacing = Pt(14)
        # Indents (no manual tabs allowed)
        style.paragraph_format.first_line_indent = Inches(0.5)

        h1 = doc.styles['Heading 1']
        h1.font.name = 'Garamond'
        h1.font.size = Pt(24)
        h1.font.bold = True
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.space_before = Pt(48)
        h1.paragraph_format.space_after = Pt(24)

        # Apply 6x9 margins with dynamic gutter
        for section in doc.sections:
            section.page_width = Inches(6)
            section.page_height = Inches(9)
            
            # Outside margins (Top, Bottom, Right/Outside)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            
            # Inside Margin (Gutter setup) docx handles it as left_margin in single view setup, 
            # or mirror margins if set, but we use left margin as basic gutter offset here
            section.left_margin = Inches(0.5 + gutter_inches) 

        title = self.metadata.get("title", "Untitled Manuscript")
        author = self.metadata.get("author_name", "Unknown Author")

        # Title Page Generator (no headers/footers ideally, but simple for now)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Inches(2)
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(36)

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(f"by {author}")
        run.font.size = Pt(18)

        doc.add_page_break()

        # Build Copyright Page
        copyright_text = f"""Copyright © {datetime.now().year} by {author}

All rights reserved. No part of this publication may be reproduced, distributed, or transmitted in any form or by any means, including photocopying, recording, or other electronic or mechanical methods, without the prior written permission of the publisher.

This is a work of fiction. Names, characters, businesses, places, events, locales, and incidents are either the products of the author's imagination or used in a fictitious manner.

First Edition: {datetime.now().strftime("%B %Y")}
Published by AI Book Factory
"""
        copy_para = doc.add_paragraph(copyright_text)
        copy_para.style = doc.styles['Normal']
        copy_para.paragraph_format.first_line_indent = Inches(0)
        copy_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Table of Contents logic (simplified clickable list)
        doc.add_heading('Table of Contents', level=1)
        
        # We need the outline to build the TOC, let's parse titles first
        with open(txt_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        
        chapter_titles = []
        for line in lines:
            if line.strip().startswith("CHAPTER "):
                title_text = line.strip().split(":", 1)[-1].strip() if ":" in line else line.strip()
                chapter_titles.append(title_text)

        for i, ct in enumerate(chapter_titles):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            
            # Internal link logic: Target must match the bookmark name
            # Note: docx internal hyperlinking is non-trivial, using standard text for now
            # but ensuring the structure is ready for the heading styles that Word auto-links.
            p.add_run(f"{i+1}. {ct}")
            
        doc.add_page_break()

        # Read the raw text and parse
        with open(txt_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        # Start content chunking
        in_chapter = False
        
        for idx, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue
                
            if stripped.startswith("CHAPTER"):
                doc.add_page_break()
                in_chapter = True
                
                # Excerpt "CHAPTER 1: Title"
                title_text = stripped.split(":", 1)[-1].strip() if ":" in stripped else stripped
                h_para = doc.add_heading(title_text, level=1)
                
                # Add Bookmark for TOC (matching the name used in Word or TOC logic)
                # self._add_bookmark(h_para, f"chapter_{i}") # Simplified bookmarking
                
                # Insert Chapter Illustration if exists
                art_count = sum(1 for i in range(idx) if lines[i].strip().startswith("CHAPTER ")) + 1
                art_path = self.book_dir / f"chapter_{art_count}.png"
                if art_path.exists():
                    doc.add_picture(str(art_path), width=Inches(4.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("\n") # Spacer
                
            elif stripped.startswith("======"):
                continue
                
            elif in_chapter:
                # Add a normal content paragraph 
                para = doc.add_paragraph(stripped)
                para.style = doc.styles['Normal']
                
        output_path = self.book_dir / "book_kdp.docx"
        doc.save(str(output_path))
        print(f"✅ Generated compliant DOCX: {output_path}")
        return output_path

def format_all_books():
    books_dir = Path("books")
    if not books_dir.exists():
        print("No books directory found to format.")
        return
        
    for item in books_dir.iterdir():
        if item.is_dir():
            print(f"Formatting {item.name}...")
            formatter = KDPFormatter(item)
            try:
                formatter.format_for_kdp()
            except Exception as e:
                print(f"Error formatting {item.name}: {e}")

if __name__ == "__main__":
    format_all_books()
